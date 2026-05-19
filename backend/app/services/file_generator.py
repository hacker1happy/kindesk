import os
import re
from docx import Document

NON_CLAIMANT_NOTE = " (NON CLAIMANT NOC GIVEN)"
NON_CLAIMANT_LEGAL_HEIR_DOCS = {
    "1. Authorization_Letter.docx",
    "2. Request_Letter.docx",
    "3. ISR1.docx",
    "4. SH-13.docx",
    "5. ISR4.docx",
    "6. FormA.docx",
    "7. FormB_Indemnity.docx",
    "8. ISR5-AnnexureC.docx",
}
FORM_A_B_NON_CLAIMANT_LINES = {
    "B": "I/We, LEGALHEIRB son / daughter of LHBFATHER residence of LHBADDRESS having Permanent Account No (s) LHBPAN",
    "C": "I/We, LEGALHEIRC son / daughter of LHCFATHER residence of LHCADDRESS having Permanent Account No (s) LHCPAN",
}
PRESERVE_CASE_KEYS = {"REQUESTLETTERSUBJECT"}
LEGAL_HEIR_KEYS = {"LEGALHEIRA", "LEGALHEIRB", "LEGALHEIRC"}
PLACEHOLDER_PATTERN = re.compile(r"[A-Za-z][A-Za-z0-9]*")


def replacement_value(key, value):
    if key in PRESERVE_CASE_KEYS:
        return str(value)

    return str(value).upper()


def replace_text(paragraph, key, value):
    if key in paragraph.text:
        for run in paragraph.runs:
            if key in run.text:
                run.text = run.text.replace(key, replacement_value(key, value))


def replace_legal_heir_with_non_claimant_note(paragraph, key, value):
    if key not in paragraph.text:
        return

    for run in paragraph.runs:
        if key not in run.text:
            continue

        before, after = run.text.split(key, 1)
        run.text = before + replacement_value(key, value)

        note_run = paragraph.add_run(NON_CLAIMANT_NOTE)
        note_run.bold = True
        note_run.underline = True

        if after:
            paragraph.add_run(after)


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def remove_non_claimant_form_lines(doc, file_name, non_claimant_suffixes):
    if file_name not in {"6. FormA.docx", "7. FormB_Indemnity.docx"}:
        return

    lines_to_remove = [
        FORM_A_B_NON_CLAIMANT_LINES[suffix]
        for suffix in non_claimant_suffixes
        if suffix in FORM_A_B_NON_CLAIMANT_LINES
    ]
    if not lines_to_remove:
        return

    for paragraph in list(iter_paragraphs(doc)):
        paragraph_text = " ".join(paragraph.text.split())
        if any(" ".join(line.split()) in paragraph_text for line in lines_to_remove):
            remove_paragraph(paragraph)


def replace_document_text(doc, data, file_name):
    non_claimant_suffixes = data.get("_NON_CLAIMANT_SUFFIXES", [])
    can_show_non_claimant_note = file_name in NON_CLAIMANT_LEGAL_HEIR_DOCS
    replacements = {key: value for key, value in data.items() if not key.startswith("_")}
    paragraphs = list(iter_paragraphs(doc))

    for paragraph in paragraphs:
        paragraph_text = paragraph.text
        if not paragraph_text:
            continue

        keys_in_paragraph = sorted(
            {key for key in PLACEHOLDER_PATTERN.findall(paragraph_text) if key in replacements},
            key=len,
            reverse=True,
        )

        for key in keys_in_paragraph:
            value = replacements[key]
            suffix = key.replace("LEGALHEIR", "", 1)
            if (
                can_show_non_claimant_note
                and key in LEGAL_HEIR_KEYS
                and suffix in non_claimant_suffixes
            ):
                replace_legal_heir_with_non_claimant_note(paragraph, key, value)
            else:
                replace_text(paragraph, key, value)

            paragraph_text = paragraph.text


def generate_documents(files, output_dir, data):

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for template in files:
        doc = Document(template)
        file_name = os.path.basename(template)

        remove_non_claimant_form_lines(doc, file_name, data.get("_NON_CLAIMANT_SUFFIXES", []))
        replace_document_text(doc, data, file_name)

        output_path = os.path.join(output_dir, os.path.basename(template))
        doc.save(output_path)

    return output_dir
